import tkinter as tk
import math
from datetime import datetime
from tkinter import filedialog
from docx import Document
from classes.Document import MyDocument
import os
from Controller.Controller import Controller

documents = []
id_ = 0
def open_word_file(id_, documents):
    file_path = filedialog.askopenfilenames(filetypes=[("Word Files", "*.docx")])
    if file_path:
        for path in file_path:
            doc = Document(path)
            doc_name = os.path.basename(path)
            doc_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            doc_created_date = datetime.fromtimestamp(os.path.getctime(path)).strftime('%H:%M - %d.%m.%Y').split("-")
            #doc_modified_date = datetime.fromtimestamp(os.path.getmtime(path)).strftime('%H:%M - %d.%m.%Y').split("-")
            document = MyDocument(doc_name,doc_content,doc_created_date[1],doc_created_date[0],id_)
            id_+=1

            document.add_document_to_base(documents)
        print(documents)

def calculate_idf(documents):
    # Создайте словарь для хранения числа документов, содержащих каждый термин
    term_document_count = {}

    # Общее количество документов
    total_documents = len(documents)
    termins = []
    # Проход по каждому документу
    for doc in documents:
        # Получите уникальные термины в текущем документе
        unique_terms = set(doc.text.split())

        # Увеличьте счетчик для каждого термина
        for term in unique_terms:
            termins.append(term)
    for doc in documents:
        for term in termins:
            if term in set(doc.text.split()):
                term_document_count[term] = term_document_count.get(term, 0) + 1

    # Рассчитайте IDF для каждого термина
    idf_values = {}
    for term, doc_count in term_document_count.items():
        idf = math.log(total_documents / (doc_count + 1))  # Добавляем 1, чтобы избежать деления на 0
        idf_values[term] = idf

    print(idf_values)
    return idf_values



if __name__ == "__main__":
    root = tk.Tk()
    root.geometry("900x600")
    app = Controller(root)
    root.mainloop()

# # Создаем главное окно
# root = tk.Tk()
# root.title("Выбор Word-файла")
#
# # Создаем кнопку для выбора файла
# open_button = tk.Button(root, text="Добавить Word-документ", command=lambda : open_word_file(id_, documents))
# open_button.pack()
#
# open_button = tk.Button(root, text="ИДФ", command=lambda : calculate_idf(documents))
# open_button.pack()
#
# # Создаем метку для вывода информации о файле
# result_label = tk.Label(root, text="")
# result_label.pack()
#
# root.mainloop()
