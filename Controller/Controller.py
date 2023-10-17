import nltk
from Model.Model import Model
from View.View import View
from docx import Document
from datetime import datetime
from classes.Document import MyDocument
import os
import string
import pymorphy2
from tkinter import filedialog
from tkinter import messagebox
import tkinter as tk
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import math
import numpy as np
import re
import heapq



class Controller:
    def __init__(self, root):
        self.model = Model()
        self.view = View(root,self)

    def __punctuation(self, str):
        punctuation = string.punctuation
        translator = str.maketrans('', '', punctuation)
        result = str.translate(translator)
        characters_to_remove = ['"', '“', '”', '«', '»']

        for char in characters_to_remove:
            result = result.replace(char, '')
        return  result

    def __get_synonyms(self,word):
        morph = pymorphy2.MorphAnalyzer()
        normal_form = morph.parse(word)[0].normal_form
        synonyms = []
        for synset in morph.parse(normal_form)[0].lexeme:
            synonyms.append(synset.word)
        return synonyms


    def create_dictionary_by_documents(self):
        dictionary = []
        documents = self.model.get_documents()
        for doc in documents:
            dictionary+=self.__punctuation(doc.text.lower()).split()
        dictionary = list(set(dictionary))
        self.model.set_dictionary(dictionary)

    def create_binary_vector_documents(self):
        dictionary = self.model.get_dictionary()
        docs = self.model.get_documents()
        matrix_of_docs = []

        for doc in docs:
            vector_doc = []
            for word in dictionary:
                vector_doc.append(1 if word in doc.text else 0)
            matrix_of_docs.append(vector_doc)

        self.model.set_docs_vectors(matrix_of_docs)


    def create_binary_vector_query(self, query):
        query = self.__punctuation(query).lower()
        query = query.split()
        query_termins_synonyms = []

        for word in query:
            query_termins_synonyms+= list(set(self.__get_synonyms(word)))

        dictionary = self.model.get_dictionary()
        vector_binary_query = []

        for word in dictionary:
                vector_binary_query.append(1 if word in query_termins_synonyms else 0)

        self.model.set_query_vector(vector_binary_query)

    def calculate_similar(self):

        matrix_docs = self.model.get_docs_vectors()
        query_vector = np.array(self.model.get_query_vector())
        e_query_vector = np.linalg.norm(query_vector)
        similar = {}
        id = 0

        for vector in matrix_docs:
            vec = np.array(vector)
            e_vec = np.linalg.norm(vec)

            if (e_vec * e_query_vector) != 0:
                query_equals_doc = (np.dot(vec, query_vector))/(e_vec * e_query_vector)
                similar[id]=query_equals_doc
                id+=1

            else:
                query_equals_doc = "Nan"
                similar[id] = query_equals_doc
                id += 1

        sorted_similar = {k: v for k, v in sorted(similar.items(),reverse=True, key=lambda item: item[1])}

        self.model.set_result_similar(sorted_similar)


    def open_word_file(self):

        documents = []
        file_path = filedialog.askopenfilenames(filetypes=[("Word Files", "*.docx")])

        if file_path:
            for path in file_path:

                doc = Document(path)
                doc_name = os.path.basename(path)
                doc_content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                doc_created_date = datetime.fromtimestamp(os.path.getctime(path)).strftime('%H:%M - %d.%m.%Y').split(
                    "-")
                document = MyDocument(doc_name, path, doc_content, doc_created_date[1], doc_created_date[0])
                documents.append(document)

            self.model.set_documents(documents)
            self.update_log("Files uploaded")

    def update_log(self, message):
        self.view.log_text.config(state=tk.NORMAL)  # Делаем текстовое поле активным
        self.view.log_text.insert(tk.END, message + "\n")  # Добавляем запись
        self.view.log_text.config(state=tk.DISABLED)  # Делаем текстовое поле неактивным
        self.view.log_text.see(tk.END)

    def check_is_nan(self, similar):
        for key, value in similar.items():
            if value == "Nan":
                self.update_log("Совпадения не найдены.")
                return False
            else: return True

    def start(self):
        if not self.model.get_documents():
            messagebox.showinfo("Ошибка", "Вы не загрузили документы")
            return 0

        if not self.view.query_entry.get():
            messagebox.showinfo("Ошибка", "Введите языковой запрос")
            return 0

        self.create_dictionary_by_documents()
        self.create_binary_vector_documents()
        self.create_binary_vector_query(self.view.query_entry.get())
        self.calculate_similar()

        if not self.check_is_nan(self.model.get_result_similar()):
            return 0

        docs_id = list(self.model.get_result_similar().keys())

        self.update_log("Наиболее подходящие документы:")
        for id in range(len(docs_id)):
            self.update_log(f"{id+1}. "+self.model.get_document_by_id(docs_id[id]).title + f": {self.model.get_result_similar()[docs_id[id]]}")

        self.view.show_open_files_button()

    def generate_annotation(self):
        path = f"../docs/"
        article_text = ""
        selected_index = self.view.listbox.curselection()

        if selected_index:
            selected_file = self.view.listbox.get(selected_index[0])
            file_path = os.path.join(path, selected_file)
            print(file_path)
            try:
                if file_path.endswith('.docx'):
                    doc = Document(file_path)
                    for paragraph in doc.paragraphs:
                        article_text += paragraph.text + '\n'
                elif file_path.endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8') as file:
                        article_text = file.read()
                else:
                    print("Неподдерживаемый формат файла.")
            except Exception as e:
                print(f"Произошла ошибка при чтении файла: {e}")

        print(article_text)

        article_text = re.sub(r'\[[0-9]*\]', ' ', article_text)
        article_text = re.sub(r'\s+', ' ', article_text)


        formatted_article_text = re.sub('[^а-яА-Я]', ' ', article_text)
        formatted_article_text = re.sub(r'\s+', ' ', formatted_article_text)

        sentence_list = nltk.sent_tokenize(article_text)

        stopwords = nltk.corpus.stopwords.words('russian')

        word_frequencies = {}
        for word in nltk.word_tokenize(formatted_article_text):
            if word not in stopwords:
                if word not in word_frequencies.keys():
                    word_frequencies[word] = 1
                else:
                    word_frequencies[word] += 1
        print(word_frequencies.values())
        maximum_frequency = max(word_frequencies.values())

        for word in word_frequencies.keys():
            word_frequencies[word] = (word_frequencies[word] / maximum_frequency)

        sentence_scores = {}
        for sent in sentence_list:
            for word in nltk.word_tokenize(sent.lower()):
                if word in word_frequencies.keys():
                    if len(sent.split(' ')) < 30:
                        if sent not in sentence_scores.keys():
                            sentence_scores[sent] = word_frequencies[word]
                        else:
                            sentence_scores[sent] += word_frequencies[word]

        summary_sentences = heapq.nlargest(3, sentence_scores, key=sentence_scores.get)

        summary = ' '.join(summary_sentences)

        self.update_log(f"\n{selected_file}: {summary}")

    def update_file_list(self):
        docs_id = list(self.model.get_result_similar().keys())
        self.view.listbox.delete(0, tk.END)
        for id in range(len(docs_id)):
            self.view.listbox.insert(tk.END, self.model.get_document_by_id(docs_id[id]).title)

    def open_new_files(self):
        path = f"../docs/"
        selected_index = self.view.listbox.curselection()
        if selected_index:
            selected_file = self.view.listbox.get(selected_index[0])
            os.startfile(path+selected_file)

    def recall_metric(self, a, c): # and average precision
        return a/(a+c) #r

    def precision_metric(self, a, b):
        return a/(a+b) # p

    def accuracy_metric(self, a, b, c, d):
        return (a+d)/(a+b+c+d)

    def error_metric(self, a, b, c, d):
        return (b+c)/(a+b+c+d)

    def f_measure_metric(self,r, p):
        return 2/((1/p)+(1/r))

    def precision_n_metric(self,a):
        return a/3

    def r_precision_metric(self, a):
        return 2/a


    def grafik(self):
        recall = np.array([0.1, 0.2, 0.3, 0.4, 0.5, 0.6, 0.7, 0.8, 0.85, 0.9, 1.0])
        p = []
        for el in recall:
            if el > 0.5:
             p.append(0)
            else: p.append(1)
        p = np.array(p)
        # Сортируем оценки уверенности в порядке убывания
        sorted_indices = np.argsort(recall)[::-1]
        p_sorted = p[sorted_indices]

        # Инициализируем списки для хранения точности и полноты на 11 уровнях
        precision_at_recall = []
        recall_levels = np.linspace(0, 1, 11)  # 11 равномерно распределенных уровней полноты от 0 до 1

        # Вычисляем точность на каждом уровне полноты
        for recall_level in recall_levels:
            cutoff = int(recall_level * len(p_sorted))
            y_true_cutoff = p_sorted[:cutoff]
            precision = np.sum(y_true_cutoff) / (
                        cutoff + 1e-9)  # Добавляем маленькое значение для избегания деления на ноль
            precision_at_recall.append(precision)

        # Интерполируем значения
        interpolated_precision = np.maximum.accumulate(precision_at_recall[::-1])[::-1]

        # Создаем фигуру matplotlib
        fig = Figure(figsize=(8, 6))
        ax = fig.add_subplot(111)

        # Построение кривой полноты/точности с точками и интерполированными значениями
        ax.step(recall_levels, precision_at_recall, marker='o', label='Точки')
        ax.plot(recall_levels, interpolated_precision, linestyle='--', label='Интерполированная линия')
        ax.set_xlabel('Полнота (Recall)')
        ax.set_ylabel('Точность (Precision)')
        ax.set_title('Кривая полноты/точности с интерполированными значениями')
        ax.grid(True)
        ax.legend()

        canvas = FigureCanvasTkAgg(fig, master=self.view.metrics_window)
        canvas_widget = canvas.get_tk_widget()
        canvas_widget.pack()

    def calculate_metrics(self):
        amount_relevant_docs = len(self.model.get_relevant_documents()) # a
        amount_irrelevant_docs = len(self.model.get_irrelevant_documents()) # b
        amount_bad_relevant_docs = len(self.model.get_bad_relevant_documents()) # d
        not_finded_docs = 0 # c

        reccal = self.recall_metric(amount_relevant_docs,not_finded_docs)
        precision = self.precision_metric(amount_relevant_docs, amount_irrelevant_docs)

        accuracy = self.accuracy_metric(amount_relevant_docs,amount_irrelevant_docs, not_finded_docs, amount_bad_relevant_docs)
        error = self.error_metric(amount_relevant_docs,amount_irrelevant_docs, not_finded_docs, amount_bad_relevant_docs)

        f_measure = self.f_measure_metric(reccal, precision)
        precision_n = self.precision_n_metric(amount_relevant_docs)

        r_precision = self.r_precision_metric(amount_relevant_docs)

        txt = f"Recall: {reccal} \n" \
              f"Precision: {precision}\n" \
              f"Average precision: {reccal}\n" \
              f"Accuracy: {accuracy}\n" \
              f"F-measure: {f_measure}\n" \
              f"Precision by n: {[precision_n]}\n" \
              f"R-precision: {r_precision}\n"
        self.view.label_metrics.config(text=txt)






    def calculate_idfs(self):
        # Создайте словарь для хранения числа документов, содержащих каждый термин
        term_document_count = {}
        documents = self.model.get_documents

        total_documents = len(documents)
        termins = []

        for doc in documents:

            unique_terms = set(doc.text.split())

            for term in unique_terms:
                termins.append(term)

        for doc in documents:
            for term in termins:
                if term in set(doc.text.split()):
                    term_document_count[term] = term_document_count.get(term, 0) + 1


        idf_values = {}
        for term, doc_count in term_document_count.items():
            idf = math.log(total_documents / (doc_count + 1))  # Добавляем 1, чтобы избежать деления на 0
            idf_values[term] = idf


        self.model.set_IDFS(idf_values)

    def calculated_weight_termins_and_L_vector_in_documents(self):
        documents = self.model.get_documents()
        IDFS = self.model.get_IDFS()
        WTDS = []
        L_vector = []

        if not IDFS:
            return False

        for doc in documents:
            term_document_count = {}
            Li = []

            for key in IDFS:
                term_document_count[key] = doc.text.count(key) * IDFS[key]

                if key in doc.text:
                    Li.append(1)
                else:
                    Li.append(0)

            L_vector.append(Li)
            WTDS.append(term_document_count)

        self.model.set_L_vector(L_vector)
        self.model.set_WTDS(WTDS)

    def search_query_transformation(self, user_query):
        user_termins = set(user_query.split())
        IDFS = self.model.get_IDFS()
        query_vector = []

        for termin in user_termins:
            if termin in IDFS:
                value = IDFS[termin] * user_query.count(termin)
                query_vector.append(value)

        self.model.set_query_vector(query_vector)


